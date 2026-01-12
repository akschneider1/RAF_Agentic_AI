"""
Document Processing - Parse and structure legal/regulatory documents
"""

import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Iterator, Optional


class DocumentType(Enum):
    """Types of government documents"""
    STATUTE = "statute"
    REGULATION = "regulation"
    ADMINISTRATIVE_CODE = "administrative_code"
    POLICY = "policy"
    GUIDANCE = "guidance"
    EXECUTIVE_ORDER = "executive_order"
    BILL = "bill"
    UNKNOWN = "unknown"


@dataclass
class Citation:
    """A legal citation reference"""
    raw_text: str
    title: Optional[str] = None
    chapter: Optional[str] = None
    section: Optional[str] = None
    subsection: Optional[str] = None
    paragraph: Optional[str] = None

    def __str__(self) -> str:
        parts = []
        if self.title:
            parts.append(f"Title {self.title}")
        if self.chapter:
            parts.append(f"Ch. {self.chapter}")
        if self.section:
            parts.append(f"§ {self.section}")
        if self.subsection:
            parts.append(f"({self.subsection})")
        return " ".join(parts) if parts else self.raw_text


@dataclass
class Section:
    """A section of a legal document"""
    id: str
    title: Optional[str]
    text: str
    citation: Optional[Citation] = None
    parent_id: Optional[str] = None
    children: list["Section"] = field(default_factory=list)
    cross_references: list[Citation] = field(default_factory=list)
    effective_date: Optional[str] = None
    last_amended: Optional[str] = None

    @property
    def word_count(self) -> int:
        return len(self.text.split())

    @property
    def sentence_count(self) -> int:
        return len(re.findall(r"[.!?]+", self.text))


@dataclass
class LegalDocument:
    """A parsed legal document (statute, regulation, policy, etc.)"""
    title: str
    doc_type: DocumentType
    source: str
    sections: list[Section] = field(default_factory=list)
    effective_date: Optional[str] = None
    jurisdiction: Optional[str] = None
    raw_text: Optional[str] = None

    @property
    def total_sections(self) -> int:
        return len(self.sections)

    @property
    def total_words(self) -> int:
        return sum(s.word_count for s in self.sections)

    def iter_sections(self) -> Iterator[Section]:
        """Iterate through all sections including nested ones"""
        def _iter(sections: list[Section]) -> Iterator[Section]:
            for section in sections:
                yield section
                yield from _iter(section.children)
        yield from _iter(self.sections)


class DocumentProcessor:
    """
    Process and parse legal/regulatory documents into structured format.

    Handles various document formats and extracts:
    - Section hierarchy
    - Citations and cross-references
    - Definitions
    - Effective dates
    """

    # Common patterns in legal text
    SECTION_PATTERNS = [
        r"(?:Section|Sec\.|§)\s*(\d+[\w.-]*)",
        r"(?:Article)\s+([IVXLCDM]+|\d+)",
        r"(?:Chapter|Ch\.)\s+(\d+[\w.-]*)",
        r"(?:Part)\s+([A-Z]|\d+)",
    ]

    CITATION_PATTERN = re.compile(
        r"(?:(?P<title>\d+)\s+)?(?:U\.?S\.?C\.?|C\.?F\.?R\.?|Stat\.?)\s*"
        r"(?:§|Section)?\s*(?P<section>[\d.]+)"
        r"(?:\s*\((?P<subsection>[a-z0-9]+)\))?"
    )

    CROSS_REF_PATTERN = re.compile(
        r"(?:pursuant to|under|as defined in|see|refer to)\s+"
        r"(?:section|§)\s*([\d.]+(?:\([a-z]\))?)",
        re.IGNORECASE
    )

    DEFINITION_PATTERN = re.compile(
        r'"([^"]+)"\s+means\s+([^.]+\.)',
        re.IGNORECASE
    )

    def __init__(self):
        self.documents: list[LegalDocument] = []

    def parse_text(
        self,
        text: str,
        doc_type: DocumentType = DocumentType.UNKNOWN,
        title: str = "Untitled",
        source: str = "unknown",
        jurisdiction: Optional[str] = None,
    ) -> LegalDocument:
        """Parse raw text into a structured LegalDocument"""

        sections = self._extract_sections(text)
        cross_refs = self._extract_cross_references(text)

        # Attach cross-references to relevant sections
        for section in sections:
            section.cross_references = [
                ref for ref in cross_refs
                if ref.raw_text in section.text
            ]

        doc = LegalDocument(
            title=title,
            doc_type=doc_type,
            source=source,
            sections=sections,
            jurisdiction=jurisdiction,
            raw_text=text,
        )

        self.documents.append(doc)
        return doc

    def _extract_sections(self, text: str) -> list[Section]:
        """Extract sections from document text"""
        sections = []

        # Try to find section boundaries
        section_starts = []
        for pattern in self.SECTION_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                section_starts.append((match.start(), match.group(0), match.group(1)))

        # Sort by position
        section_starts.sort(key=lambda x: x[0])

        if not section_starts:
            # No sections found, treat entire text as one section
            sections.append(Section(
                id="1",
                title=None,
                text=text.strip(),
            ))
            return sections

        # Extract text between section markers
        for i, (start, header, section_id) in enumerate(section_starts):
            # Find end of this section
            if i + 1 < len(section_starts):
                end = section_starts[i + 1][0]
            else:
                end = len(text)

            section_text = text[start:end].strip()

            # Try to extract title from first line
            lines = section_text.split("\n", 2)
            title = None
            if len(lines) > 1 and len(lines[0]) < 200:
                title = lines[0].strip()
                section_text = "\n".join(lines[1:]).strip()

            sections.append(Section(
                id=section_id,
                title=title,
                text=section_text,
                citation=Citation(raw_text=header, section=section_id),
            ))

        return sections

    def _extract_cross_references(self, text: str) -> list[Citation]:
        """Extract cross-references to other sections/statutes"""
        refs = []

        for match in self.CROSS_REF_PATTERN.finditer(text):
            refs.append(Citation(
                raw_text=match.group(0),
                section=match.group(1),
            ))

        return refs

    def extract_definitions(self, text: str) -> dict[str, str]:
        """Extract defined terms from legal text"""
        definitions = {}

        for match in self.DEFINITION_PATTERN.finditer(text):
            term = match.group(1).lower()
            definition = match.group(2).strip()
            definitions[term] = definition

        return definitions

    def detect_document_type(self, text: str) -> DocumentType:
        """Attempt to detect the type of document from its content"""

        text_lower = text.lower()

        if "administrative code" in text_lower or "admin. code" in text_lower:
            return DocumentType.ADMINISTRATIVE_CODE
        elif "c.f.r." in text_lower or "code of federal regulations" in text_lower:
            return DocumentType.REGULATION
        elif "u.s.c." in text_lower or "united states code" in text_lower:
            return DocumentType.STATUTE
        elif "executive order" in text_lower:
            return DocumentType.EXECUTIVE_ORDER
        elif any(x in text_lower for x in ["bill", "be it enacted", "a bill to"]):
            return DocumentType.BILL
        elif "policy" in text_lower or "guidance" in text_lower:
            return DocumentType.POLICY

        return DocumentType.UNKNOWN

    def compute_complexity_score(self, section: Section) -> float:
        """
        Compute a readability/complexity score for a section.

        Higher score = more complex/harder to understand
        """
        text = section.text

        # Factors that increase complexity
        avg_sentence_length = section.word_count / max(section.sentence_count, 1)
        cross_ref_count = len(section.cross_references)
        nested_clauses = len(re.findall(r"\([a-z]\)", text))
        legal_jargon = len(re.findall(
            r"\b(herein|thereof|whereby|notwithstanding|pursuant|aforementioned)\b",
            text, re.IGNORECASE
        ))
        passive_voice = len(re.findall(r"\b(shall be|is|are|was|were)\s+\w+ed\b", text))

        # Weighted score
        score = (
            (avg_sentence_length / 20) * 0.3 +  # Normalize to ~1 for 20-word sentences
            (cross_ref_count / 5) * 0.2 +
            (nested_clauses / 10) * 0.2 +
            (legal_jargon / 5) * 0.15 +
            (passive_voice / 10) * 0.15
        )

        return min(score, 1.0)  # Cap at 1.0

    def parse_file(self, filepath: Path) -> LegalDocument:
        """Parse a document file (supports .txt, .pdf, .docx)"""
        suffix = filepath.suffix.lower()

        if suffix == ".txt":
            text = filepath.read_text()
        elif suffix == ".pdf":
            text = self._parse_pdf(filepath)
        elif suffix == ".docx":
            text = self._parse_docx(filepath)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        doc_type = self.detect_document_type(text)

        return self.parse_text(
            text=text,
            doc_type=doc_type,
            title=filepath.stem,
            source=str(filepath),
        )

    def _parse_pdf(self, filepath: Path) -> str:
        """Extract text from PDF"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            return "\n".join(page.extract_text() for page in reader.pages)
        except ImportError:
            raise ImportError("pypdf required for PDF parsing: pip install pypdf")

    def _parse_docx(self, filepath: Path) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join(para.text for para in doc.paragraphs)
        except ImportError:
            raise ImportError("python-docx required for DOCX parsing: pip install python-docx")
