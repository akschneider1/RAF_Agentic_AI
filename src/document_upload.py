"""
Document upload and processing module for RAF Agentic AI.

Supports:
- PDF text extraction
- DOCX/DOC file processing
- URL scraping for online legal codes
- Text file upload
"""

import io
import re
from typing import Optional
from dataclasses import dataclass
from enum import Enum


class DocumentFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    URL = "url"
    UNKNOWN = "unknown"


@dataclass
class ExtractedDocument:
    """Result of document extraction"""
    text: str
    title: str
    source: str
    format: DocumentFormat
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    extraction_notes: Optional[str] = None

    def __post_init__(self):
        if self.word_count is None:
            self.word_count = len(self.text.split())


def detect_format(filename: str) -> DocumentFormat:
    """Detect document format from filename"""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return DocumentFormat.PDF
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        return DocumentFormat.DOCX
    elif filename_lower.endswith('.txt'):
        return DocumentFormat.TXT
    elif filename_lower.startswith('http'):
        return DocumentFormat.URL
    return DocumentFormat.UNKNOWN


def extract_text_from_pdf(file_bytes: bytes, filename: str = "document.pdf") -> ExtractedDocument:
    """Extract text from PDF file"""
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        full_text = "\n\n".join(pages)

        # Try to extract title from metadata or first line
        title = filename
        if reader.metadata and reader.metadata.title:
            title = reader.metadata.title
        elif pages:
            first_line = pages[0].split('\n')[0][:100]
            if first_line:
                title = first_line

        return ExtractedDocument(
            text=full_text,
            title=title,
            source=filename,
            format=DocumentFormat.PDF,
            page_count=len(reader.pages),
        )
    except ImportError:
        return ExtractedDocument(
            text="",
            title=filename,
            source=filename,
            format=DocumentFormat.PDF,
            extraction_notes="PDF extraction requires pypdf library. Install with: pip install pypdf"
        )
    except Exception as e:
        return ExtractedDocument(
            text="",
            title=filename,
            source=filename,
            format=DocumentFormat.PDF,
            extraction_notes=f"PDF extraction failed: {str(e)}"
        )


def extract_text_from_docx(file_bytes: bytes, filename: str = "document.docx") -> ExtractedDocument:
    """Extract text from DOCX file"""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Try to get title from core properties
        title = filename
        if doc.core_properties.title:
            title = doc.core_properties.title
        elif paragraphs:
            title = paragraphs[0][:100]

        return ExtractedDocument(
            text=full_text,
            title=title,
            source=filename,
            format=DocumentFormat.DOCX,
        )
    except ImportError:
        return ExtractedDocument(
            text="",
            title=filename,
            source=filename,
            format=DocumentFormat.DOCX,
            extraction_notes="DOCX extraction requires python-docx library. Install with: pip install python-docx"
        )
    except Exception as e:
        return ExtractedDocument(
            text="",
            title=filename,
            source=filename,
            format=DocumentFormat.DOCX,
            extraction_notes=f"DOCX extraction failed: {str(e)}"
        )


def extract_text_from_txt(file_bytes: bytes, filename: str = "document.txt") -> ExtractedDocument:
    """Extract text from plain text file"""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = file_bytes.decode('utf-8', errors='ignore')

        title = filename
        lines = text.split('\n')
        if lines:
            title = lines[0][:100] or filename

        return ExtractedDocument(
            text=text,
            title=title,
            source=filename,
            format=DocumentFormat.TXT,
        )
    except Exception as e:
        return ExtractedDocument(
            text="",
            title=filename,
            source=filename,
            format=DocumentFormat.TXT,
            extraction_notes=f"Text extraction failed: {str(e)}"
        )


def extract_text_from_url(url: str) -> ExtractedDocument:
    """Extract text from a URL (basic web scraping)"""
    try:
        import urllib.request
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {'script', 'style', 'nav', 'footer', 'header'}
                self.current_tag = None
                self.title = ""

            def handle_starttag(self, tag, attrs):
                self.current_tag = tag

            def handle_endtag(self, tag):
                self.current_tag = None

            def handle_data(self, data):
                if self.current_tag == 'title':
                    self.title = data.strip()
                elif self.current_tag not in self.skip_tags:
                    text = data.strip()
                    if text:
                        self.text_parts.append(text)

        # Fetch URL
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=10) as response:
            html = response.read().decode('utf-8', errors='ignore')

        # Parse HTML
        parser = TextExtractor()
        parser.feed(html)

        full_text = "\n".join(parser.text_parts)

        return ExtractedDocument(
            text=full_text,
            title=parser.title or url,
            source=url,
            format=DocumentFormat.URL,
        )
    except Exception as e:
        return ExtractedDocument(
            text="",
            title=url,
            source=url,
            format=DocumentFormat.URL,
            extraction_notes=f"URL extraction failed: {str(e)}"
        )


def extract_document(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """Main extraction function - routes to appropriate extractor"""
    doc_format = detect_format(filename)

    if doc_format == DocumentFormat.PDF:
        return extract_text_from_pdf(file_bytes, filename)
    elif doc_format == DocumentFormat.DOCX:
        return extract_text_from_docx(file_bytes, filename)
    elif doc_format == DocumentFormat.TXT:
        return extract_text_from_txt(file_bytes, filename)
    else:
        # Try as plain text
        return extract_text_from_txt(file_bytes, filename)


def clean_legal_text(text: str) -> str:
    """Clean and normalize legal text"""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)

    # Fix common OCR issues in legal text
    text = text.replace('§', 'Section ')
    text = re.sub(r'(\d+)\s*\.\s*(\d+)', r'\1.\2', text)  # Fix section numbers

    # Remove page numbers and headers (common patterns)
    text = re.sub(r'\n\d+\n', '\n', text)  # Standalone page numbers
    text = re.sub(r'Page \d+ of \d+', '', text)

    return text.strip()


def extract_sections(text: str) -> list[dict]:
    """Extract individual sections from legal text"""
    sections = []

    # Common section patterns
    patterns = [
        r'(Section\s+\d+[\.\d]*[^\n]*)\n(.*?)(?=Section\s+\d+|$)',
        r'(§\s*\d+[\.\d]*[^\n]*)\n(.*?)(?=§\s*\d+|$)',
        r'(\(\w+\)\s+[^\n]+)\n(.*?)(?=\(\w+\)|$)',
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
        if matches:
            for title, content in matches:
                sections.append({
                    'title': title.strip(),
                    'content': content.strip()[:2000],  # Limit content length
                })
            break

    # If no sections found, split by paragraphs
    if not sections:
        paragraphs = text.split('\n\n')
        for i, para in enumerate(paragraphs[:20]):  # Limit to first 20
            if para.strip():
                sections.append({
                    'title': f'Paragraph {i+1}',
                    'content': para.strip()[:2000],
                })

    return sections
